from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "se104-design" / "Báo cáo đồ án SE104 - Nhóm 2.docx"
IMAGES = ROOT / "docs" / "se104-design" / "images"
GENERATOR = ROOT / "scripts" / "generate_design_artifacts.py"


def load_generator():
    spec = importlib.util.spec_from_file_location("design_generator", GENERATOR)
    if spec is None or spec.loader is None:
        raise RuntimeError("Cannot load design generator")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def delete_from_heading(doc: Document, heading_text: str) -> None:
    start = None
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == heading_text or text.startswith("3. Mô hình hoá theo Chức năng"):
            start = idx
            break
    if start is None:
        raise RuntimeError(f"Heading not found: {heading_text}")
    for para in list(doc.paragraphs[start:]):
        delete_paragraph(para)


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor(31, 41, 55)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_bold_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(f"- {item}")


def add_code(doc: Document, code: str) -> None:
    for line in code.strip().splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(line)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "9CA3AF")


def add_image(doc: Document, path: Path, width_inches: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(75, 85, 99)


def add_dfd_section(doc: Document, gen) -> None:
    add_heading(doc, "3. Mô hình hoá theo Chức năng (DFD - Data Flow Diagram)", 3)
    doc.add_paragraph(
        "Sử dụng DFD để mô hình hóa chi tiết các nghiệp vụ cốt lõi của hệ thống PetCare. "
        "Mỗi sơ đồ được trình bày ở mức tối giản gồm người dùng, thiết bị nhập, tiến trình xử lý, "
        "thiết bị xuất và bộ nhớ phụ; các luồng dữ liệu D1-D6 được mô tả ngay dưới sơ đồ."
    )

    for idx, uc in enumerate(gen.USE_CASES, start=1):
        add_heading(doc, f"3.{idx}. {uc['name']}", 4)
        add_bold_label(doc, "Sơ đồ luồng dữ liệu:")
        add_image(doc, IMAGES / f"dfd-{uc['id'].lower()}.png", 5.8)
        add_caption(doc, f"Hình IV.3.{idx}. DFD - {uc['name']}")

        add_bold_label(doc, "⮚ Mô tả luồng dữ liệu:")
        add_bullets(doc, [f"{code}: {desc}" for code, desc in gen.dfd_flows(uc)])

        add_bold_label(doc, "⮚ Thuật toán:")
        add_bullets(doc, gen.dfd_algorithm(uc))

        add_bold_label(doc, "Mã PlantUML:")
        add_code(doc, gen.plantuml_dfd(uc))


def add_sequence_section(doc: Document, gen) -> None:
    doc.add_page_break()
    add_heading(doc, "4. Sơ đồ Tuần tự (Sequence Diagram)", 3)
    doc.add_paragraph(
        "Sequence Diagram thể hiện thứ tự tương tác giữa các đối tượng ở mức phân tích thiết kế. "
        "Các use case có liên quan được gộp thành một sơ đồ để thể hiện rõ nhánh xử lý và tránh lặp lại."
    )

    for idx, seq in enumerate(gen.SEQUENCE_GROUPS, start=1):
        add_heading(doc, f"4.{idx}. {seq['name']}", 4)
        add_image(doc, IMAGES / f"{seq['file']}.png", 6.7)
        add_caption(doc, f"Hình IV.4.{idx}. Sequence Diagram - {seq['name']}")
        add_bold_label(doc, "Mã PlantUML:")
        add_code(doc, gen.plantuml_sequence(seq))


def add_system_design(doc: Document, gen) -> None:
    doc.add_page_break()
    add_heading(doc, "V. THIẾT KẾ HỆ THỐNG", 1)

    add_heading(doc, "1. Kiến trúc hệ thống", 3)
    doc.add_paragraph(
        "Hệ thống PetCare được xây dựng theo mô hình Client-Server. Frontend React/Vite chịu trách nhiệm "
        "hiển thị giao diện và điều hướng; Backend FastAPI xử lý xác thực, nghiệp vụ, truy cập dữ liệu và "
        "tích hợp ngoài; PostgreSQL lưu trữ dữ liệu nghiệp vụ. Kiến trúc hiện tại là modular monolith theo "
        "lớp, thuận tiện triển khai bằng Docker Compose và vẫn tách rõ các module nghiệp vụ."
    )
    add_image(doc, IMAGES / "system-architecture.png", 6.5)
    add_caption(doc, "Hình V.1. Sơ đồ kiến trúc hệ thống PetCare")
    add_bold_label(doc, "Mã PlantUML:")
    add_code(doc, gen.SYSTEM_DIAGRAMS["system-architecture"]["plantuml"])

    add_heading(doc, "2. Mô tả các thành phần trong hệ thống", 3)
    add_bullets(
        doc,
        [
            "Frontend: gồm các page React, route điều hướng, AuthContext quản lý phiên đăng nhập và service client gọi HTTP API.",
            "Backend API: gồm các router auth, users, pets, products, bookings, care_logs, carts, orders, payments, reports, notifications và admin.",
            "Service layer: xử lý quy tắc nghiệp vụ như xác thực, phân quyền, kiểm tra chủ sở hữu thú cưng, kiểm tra slot đặt lịch, tính tổng giỏ hàng, tạo đơn hàng, tạo URL thanh toán và tổng hợp báo cáo.",
            "Repository/Persistence layer: truy vấn và cập nhật dữ liệu thông qua SQLAlchemy async ORM.",
            "Tích hợp ngoài: S3 lưu ảnh, GHN lấy địa chỉ và phí vận chuyển, VNPAY xử lý thanh toán trực tuyến.",
        ],
    )
    add_image(doc, IMAGES / "backend-layering.png", 6.1)
    add_caption(doc, "Hình V.2. Sơ đồ phân lớp backend")
    add_bold_label(doc, "Mã PlantUML:")
    add_code(doc, gen.SYSTEM_DIAGRAMS["backend-layering"]["plantuml"])

    add_heading(doc, "3. Thiết kế dữ liệu mức hệ thống", 3)
    doc.add_paragraph(
        "Mô hình dữ liệu tập trung vào các thực thể chính của hệ thống: người dùng, thú cưng, sản phẩm/dịch vụ, "
        "giỏ hàng, đơn hàng, lịch hẹn, nhật ký chăm sóc và thông báo."
    )
    add_image(doc, IMAGES / "data-model.png", 5.9)
    add_caption(doc, "Hình V.3. Mô hình dữ liệu mức khái niệm")
    add_bold_label(doc, "Mã PlantUML:")
    add_code(doc, gen.SYSTEM_DIAGRAMS["data-model"]["plantuml"])

    add_heading(doc, "4. Thiết kế bảo mật và vận hành", 3)
    add_bullets(
        doc,
        [
            "Xác thực bằng JWT access/refresh token; phân quyền theo vai trò user và admin.",
            "Các chức năng quản trị yêu cầu quyền admin; dữ liệu cá nhân như pets, bookings, cart, orders và notifications được lọc theo user_id.",
            "Middleware xử lý logging, lỗi tập trung và CORS cho frontend local.",
            "Docker Compose vận hành frontend, backend và PostgreSQL; Alembic quản lý migration cơ sở dữ liệu.",
        ],
    )

    add_heading(doc, "5. Ánh xạ Use Case - Module triển khai", 3)
    table = doc.add_table(rows=1, cols=4)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Use Case", "Frontend", "Backend/API", "Bảng dữ liệu chính"]):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    rows = [
        ("UC-01, UC-02", "auth/Login, auth/Register", "auth.py, auth_service.py", "users"),
        ("UC-03", "profile/Profile, admin/Users", "users.py, admin.py", "users"),
        ("UC-04", "pets/*", "pets.py, pet_service.py", "pets"),
        ("UC-05, UC-06, UC-09", "services/*, shop/*", "products.py, categories.py", "products, categories"),
        ("UC-07", "bookings/*", "bookings.py, booking_service.py", "bookings, pets, notifications"),
        ("UC-08", "care/*", "care_logs.py, care_log_service.py", "care_logs, pets"),
        ("UC-10", "cart/Cart", "carts.py, cart_service.py", "carts, cart_items"),
        ("UC-11, UC-12", "cart/Checkout, orders/*, payment/*", "orders.py, payments.py", "orders, order_items"),
        ("UC-13", "notifications/Notifications", "notifications.py", "notifications"),
        ("UC-14", "reports/Reports", "reports.py, report_service.py", "orders, bookings"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text


def main() -> None:
    gen = load_generator()
    doc = Document(DOCX_PATH)
    ensure_styles(doc)
    delete_from_heading(doc, "3. Mô hình hoá theo Chức năng (DFD – Data Flow Diagram)")
    add_dfd_section(doc, gen)
    add_sequence_section(doc, gen)
    add_system_design(doc, gen)
    doc.save(DOCX_PATH)
    print("Updated design DOCX")


if __name__ == "__main__":
    main()
